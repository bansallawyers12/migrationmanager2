#!/usr/bin/env python3
"""
Test script for LibreOffice DOCX to PDF Converter
"""

import requests
import json
import sys
import os

def test_health_check():
    """Test the health check endpoint"""
    print("🔍 Testing health check...")
    try:
        response = requests.get('http://localhost:5000/health', timeout=10)
        if response.status_code == 200:
            data = response.json()
            print(f"✅ Health check passed!")
            print(f"   Service: {data.get('service')}")
            print(f"   Version: {data.get('version')}")
            print(f"   LibreOffice available: {data.get('libreoffice_available')}")
            if data.get('libreoffice_path'):
                print(f"   LibreOffice path: {data.get('libreoffice_path')}")
            print(f"   Platform: {data.get('platform')}")
            return True
        else:
            print(f"❌ Health check failed: HTTP {response.status_code}")
            return False
    except requests.exceptions.ConnectionError:
        print("❌ Health check failed: Cannot connect to service")
        print("   Make sure the LibreOffice converter is running on http://localhost:5000")
        return False
    except Exception as e:
        print(f"❌ Health check failed: {str(e)}")
        return False

def test_service_info():
    """Test the service information endpoint"""
    print("\n📋 Testing service information...")
    try:
        response = requests.get('http://localhost:5000/', timeout=10)
        if response.status_code == 200:
            data = response.json()
            print(f"✅ Service info retrieved!")
            print(f"   Service: {data.get('service')}")
            print(f"   Description: {data.get('description')}")
            print(f"   Features: {', '.join(data.get('features', []))}")
            return True
        else:
            print(f"❌ Service info failed: HTTP {response.status_code}")
            return False
    except Exception as e:
        print(f"❌ Service info failed: {str(e)}")
        return False

def test_libreoffice_availability():
    """Test LibreOffice availability"""
    print("\n🔧 Testing LibreOffice availability...")
    try:
        response = requests.get('http://localhost:5000/test', timeout=10)
        if response.status_code == 200:
            data = response.json()
            if data.get('success'):
                print("✅ LibreOffice is available and ready!")
                print(f"   Path: {data.get('libreoffice_path')}")
                print(f"   Platform: {data.get('platform')}")
                return True
            else:
                print("❌ LibreOffice not available")
                print(f"   Error: {data.get('error')}")
                if 'installation_guide' in data:
                    guide = data['installation_guide']
                    print("   Installation guide:")
                    for platform, command in guide.items():
                        print(f"     {platform}: {command}")
                return False
        else:
            print(f"❌ LibreOffice test failed: HTTP {response.status_code}")
            return False
    except Exception as e:
        print(f"❌ LibreOffice test failed: {str(e)}")
        return False

def create_test_docx():
    """Create a simple test DOCX file"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Test Document', 0)
        
        # Add paragraph with formatting
        p = doc.add_paragraph('This is a ')
        p.add_run('bold').bold = True
        p.add_run(' and ')
        p.add_run('italic').italic = True
        p.add_run(' text with ')
        p.add_run('underline').underline = True
        p.add_run(' formatting.')
        
        # Add another paragraph
        doc.add_paragraph('This is a test document for LibreOffice conversion.')
        
        # Add table
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        
        # Fill table
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f'Cell {i+1},{j+1}'
        
        # Save document
        test_file = 'test_document.docx'
        doc.save(test_file)
        print(f"✅ Created test document: {test_file}")
        return test_file
        
    except ImportError:
        print("❌ python-docx not installed. Install with: pip install python-docx")
        return None
    except Exception as e:
        print(f"❌ Failed to create test document: {str(e)}")
        return None

def test_conversion(test_file):
    """Test DOCX to PDF conversion"""
    print(f"\n🔄 Testing conversion with {test_file}...")
    
    if not os.path.exists(test_file):
        print(f"❌ Test file {test_file} not found")
        return False
    
    try:
        with open(test_file, 'rb') as file:
            files = {'file': file}
            response = requests.post('http://localhost:5000/convert', files=files, timeout=60)
            
            if response.status_code == 200:
                data = response.json()
                if data.get('success'):
                    print("✅ Conversion successful!")
                    print(f"   Method: {data.get('method')}")
                    print(f"   Message: {data.get('message')}")
                    print(f"   PDF size: {data.get('pdf_size')} bytes")
                    
                    # Save the PDF
                    import base64
                    pdf_data = base64.b64decode(data['pdf_data'])
                    output_file = test_file.replace('.docx', '.pdf')
                    with open(output_file, 'wb') as pdf_file:
                        pdf_file.write(pdf_data)
                    print(f"   PDF saved as: {output_file}")
                    
                    return True
                else:
                    print(f"❌ Conversion failed: {data.get('error')}")
                    return False
            else:
                print(f"❌ Conversion failed: HTTP {response.status_code}")
                return False
                
    except Exception as e:
        print(f"❌ Conversion test failed: {str(e)}")
        return False

def cleanup_test_files():
    """Clean up test files"""
    test_files = ['test_document.docx', 'test_document.pdf']
    for file in test_files:
        if os.path.exists(file):
            try:
                os.remove(file)
                print(f"🧹 Cleaned up: {file}")
            except:
                pass

def main():
    """Main test function"""
    print("🧪 LibreOffice DOCX to PDF Converter Test")
    print("=" * 50)
    
    # Test health check
    if not test_health_check():
        print("\n❌ Service is not running. Please start the LibreOffice converter first:")
        print("   python libreoffice_converter.py")
        sys.exit(1)
    
    # Test service info
    test_service_info()
    
    # Test LibreOffice availability
    if not test_libreoffice_availability():
        print("\n❌ LibreOffice is not available. Please install LibreOffice first.")
        print("   Download from: https://www.libreoffice.org/download/download/")
        sys.exit(1)
    
    # Create test document
    test_file = create_test_docx()
    if not test_file:
        print("\n❌ Could not create test document. Skipping conversion test.")
        sys.exit(1)
    
    # Test conversion
    conversion_success = test_conversion(test_file)
    
    # Cleanup
    cleanup_test_files()
    
    # Summary
    print("\n" + "=" * 50)
    if conversion_success:
        print("🎉 All tests passed! LibreOffice converter is working correctly.")
    else:
        print("❌ Some tests failed. Please check the error messages above.")
    
    print("\n📖 For more information, see: LIBREOFFICE_CONVERTER_README.md")

if __name__ == '__main__':
    main()
