from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import io
import os
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import base64
import subprocess
import sys
import platform

app = Flask(__name__)

def extract_formatted_content_from_docx(docx_path):
    """Extract formatted content from DOCX file with styling"""
    try:
        doc = Document(docx_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Get paragraph formatting
                alignment = paragraph.alignment
                style_name = paragraph.style.name if paragraph.style else 'Normal'
                
                # Get runs (text with specific formatting)
                runs_data = []
                for run in paragraph.runs:
                    run_data = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_size': run.font.size.pt if run.font.size else 11,
                        'font_name': run.font.name if run.font.name else 'Arial',
                        'color': str(run.font.color.rgb) if run.font.color.rgb else None
                    }
                    runs_data.append(run_data)
                
                content.append({
                    'type': 'paragraph',
                    'text': paragraph.text,
                    'alignment': alignment,
                    'style': style_name,
                    'runs': runs_data
                })
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data:
                content.append({
                    'type': 'table',
                    'data': table_data
                })
        
        return content
        
    except Exception as e:
        raise Exception(f"Failed to extract formatted content from DOCX: {str(e)}")

def create_styled_pdf_from_content(content, output_path):
    """Create PDF with preserved formatting and styling"""
    try:
        doc = SimpleDocTemplate(output_path, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()
        story = []
        
        # Create custom styles
        custom_styles = {
            'Heading1': ParagraphStyle(
                'CustomHeading1',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=12,
                textColor=colors.black
            ),
            'Heading2': ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=10,
                textColor=colors.black
            ),
            'Normal': ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                textColor=colors.black
            )
        }
        
        for item in content:
            if item['type'] == 'paragraph':
                text = item['text']
                alignment = item['alignment']
                style_name = item['style']
                runs = item['runs']
                
                # Determine alignment
                if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    align = TA_CENTER
                elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    align = TA_RIGHT
                elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    align = TA_JUSTIFY
                else:
                    align = TA_LEFT
                
                # Determine style
                if 'heading' in style_name.lower() or 'title' in style_name.lower():
                    style = custom_styles['Heading1']
                elif 'subheading' in style_name.lower():
                    style = custom_styles['Heading2']
                else:
                    style = custom_styles['Normal']
                
                # Apply alignment
                style.alignment = align
                
                # Create paragraph with formatting
                if runs and len(runs) > 1:
                    # Complex formatting with multiple runs
                    formatted_text = ""
                    for run in runs:
                        run_text = run['text']
                        if run['bold']:
                            run_text = f"<b>{run_text}</b>"
                        if run['italic']:
                            run_text = f"<i>{run_text}</i>"
                        if run['underline']:
                            run_text = f"<u>{run_text}</u>"
                        formatted_text += run_text
                    
                    para = Paragraph(formatted_text, style)
                else:
                    # Simple paragraph
                    para = Paragraph(text, style)
                
                story.append(para)
                story.append(Spacer(1, 6))
                
            elif item['type'] == 'table':
                table_data = item['data']
                if table_data:
                    # Create table with styling
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 12))
        
        if story:
            doc.build(story)
            return True
        else:
            # Create empty PDF if no content
            c = canvas.Canvas(output_path, pagesize=A4)
            c.drawString(100, 750, "Empty document")
            c.save()
            return True
            
    except Exception as e:
        raise Exception(f"Failed to create styled PDF: {str(e)}")

def try_libreoffice_conversion(docx_path, pdf_path):
    """Try using LibreOffice for the best formatting preservation"""
    try:
        # Check if LibreOffice is available
        if platform.system() == "Windows":
            # Common LibreOffice paths on Windows
            libreoffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                r"C:\Program Files\LibreOffice*\program\soffice.exe"
            ]
        else:
            # Linux/Mac paths
            libreoffice_paths = ["libreoffice", "soffice", "/usr/bin/libreoffice"]
        
        libreoffice_cmd = None
        for path in libreoffice_paths:
            try:
                if platform.system() == "Windows":
                    if os.path.exists(path):
                        libreoffice_cmd = path
                        break
                else:
                    result = subprocess.run([path, "--version"], capture_output=True, text=True)
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        break
            except:
                continue
        
        if libreoffice_cmd:
            # Convert using LibreOffice
            cmd = [
                libreoffice_cmd,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(pdf_path),
                docx_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                # LibreOffice creates PDF with different name, rename it
                expected_pdf = docx_path.replace('.docx', '.pdf')
                if os.path.exists(expected_pdf):
                    os.rename(expected_pdf, pdf_path)
                    return True
            
        return False
        
    except Exception as e:
        print(f"LibreOffice conversion failed: {str(e)}")
        return False

def try_docx2pdf_conversion(docx_path, pdf_path):
    """Try using docx2pdf for better formatting preservation"""
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return True
    except Exception as e:
        print(f"docx2pdf conversion failed: {str(e)}")
        return False

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'service': 'python-docx-to-pdf-converter-enhanced',
        'version': '2.0.0',
        'features': ['formatting-preservation', 'styling-support', 'table-support', 'libreoffice-support']
    })

@app.route('/convert-json', methods=['POST'])
def convert_docx_to_pdf():
    """Convert DOCX to PDF with formatting preservation"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        # Validate file extension
        if not file.filename.lower().endswith(('.doc', '.docx')):
            return jsonify({'success': False, 'error': 'Invalid file format. Only .doc and .docx files are supported'}), 400
        
        # Create temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
            file.save(temp_docx.name)
            docx_path = temp_docx.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
            pdf_path = temp_pdf.name
        
        try:
            # Try LibreOffice first for best formatting preservation
            if try_libreoffice_conversion(docx_path, pdf_path):
                method = 'libreoffice-best-formatting'
            # Try docx2pdf for good formatting
            elif try_docx2pdf_conversion(docx_path, pdf_path):
                method = 'docx2pdf-enhanced'
            else:
                # Fallback to our enhanced formatting method
                content = extract_formatted_content_from_docx(docx_path)
                create_styled_pdf_from_content(content, pdf_path)
                method = 'python-enhanced-formatting'
            
            # Read PDF and convert to base64
            with open(pdf_path, 'rb') as pdf_file:
                pdf_data = pdf_file.read()
                pdf_base64 = base64.b64encode(pdf_data).decode('utf-8')
            
            # Clean up temporary files
            os.unlink(docx_path)
            os.unlink(pdf_path)
            
            return jsonify({
                'success': True,
                'filename': file.filename,
                'method': method,
                'message': 'Conversion completed successfully with maximum formatting preserved',
                'pdf_data': pdf_base64
            })
            
        except Exception as e:
            # Clean up temporary files on error
            if os.path.exists(docx_path):
                os.unlink(docx_path)
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
            raise e
            
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Conversion failed: {str(e)}'
        }), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)